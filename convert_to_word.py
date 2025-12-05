"""
Convert TECHNICAL_DOCUMENTATION.md to Word format.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def parse_markdown_to_word(md_file, docx_file):
    """Convert Markdown file to Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line.strip():
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            i += 1
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=2)
            i += 1
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
            i += 1
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:].strip(), level=4)
            i += 1
        # Code blocks
        elif line.startswith('```'):
            # Collect code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            
            # Add code block as monospace text
            if code_lines:
                code_text = ''.join(code_lines).rstrip()
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                para.style = 'No Spacing'
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
            i += 1
        # Bullet points
        elif line.strip().startswith('- '):
            text = line.strip()[2:]
            # Handle bold in bullet
            para = doc.add_paragraph(style='List Bullet')
            add_formatted_text(para, text)
            i += 1
        # Numbered lists
        elif re.match(r'^\d+\.\s+', line.strip()):
            text = re.sub(r'^\d+\.\s+', '', line.strip())
            para = doc.add_paragraph(style='List Number')
            add_formatted_text(para, text)
            i += 1
        # Regular paragraph
        else:
            para = doc.add_paragraph()
            add_formatted_text(para, line)
            i += 1
    
    doc.save(docx_file)
    print(f"✅ Word document created: {docx_file}")

def add_formatted_text(paragraph, text):
    """Add text with Markdown formatting (bold, italic, code)."""
    # Handle bold **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    md_file = 'TECHNICAL_DOCUMENTATION.md'
    docx_file = 'TECHNICAL_DOCUMENTATION.docx'
    parse_markdown_to_word(md_file, docx_file)

